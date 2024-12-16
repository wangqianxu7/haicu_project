from config import *
from typing import Dict, Union, List
from loguru import logger

import re
import os
from PlanRAG import PlanRAG
from client import ChatQwen
from PlanAgent import PlanAgent

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

from utils import load_txt
import datetime
from markdown import markdown
from bs4 import BeautifulSoup

def format_response(response):
    """
    格式化审查结果：
    1. 删除段落间多余的两个换行符，保留一个。
    2. 每行缩进。
    """
    lines = response.splitlines()
    formatted_lines = []
    for line in lines:
        if line.strip():  # 忽略空行
            formatted_lines.append("    " + line.strip())  # 每行缩进
    return "\n".join(formatted_lines)

def generate_doc(filename):

    # 创建一个空白文档
    Doc = Document()
    Doc.styles['Normal'].font.name = u'仿宋'
    Doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
    Doc.styles['Normal'].font.size = Pt(10.5)

    # 添加标题
    Doc.add_heading("AI审查-" + filename.split('.')[0], level=0)
    time_str = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
    emb_path = os.path.join(DB_PATH, filename.split('.')[0])
    target_files = os.path.join(OUT_PATH, filename.split('.')[0])
    target_path = os.path.join(target_files, f"AI智能审查报告_{time_str}.docx")

    client = ChatQwen()
    plan_rag = PlanRAG(emb_path)
    audit_agent = PlanAudit(plan_rag, client)

    filepath = "./DATA/guideline/audit_re.txt"
    filepath2 = "./DATA/guideline/audit_foundation.txt"
    audit_sug = load_txt(filepath)
    audit_foundation = load_txt(filepath2)
    test_history = []

    paragraph = Doc.add_paragraph()
    run = paragraph.add_run("报告书项目基本信息：")
    run.bold = True  # 设置加粗
    paragraph.add_run("\n" + audit_agent.base_info + "\n")

    for i in range(1, 3):
        paragraph = Doc.add_paragraph()
        run = paragraph.add_run("\n" + audit_sug[i-1])  # 注意索引从1开始，所以需要减1
        run.bold = True  # 加粗
        response, history_rs = audit_agent.chat_llm(audit_sug[i-1], audit_foundation[i-1], 1, 3, test_history)
        # response_clean = format_response(response)  # 格式化响应
        logger.info(f"审查要求{i}：{audit_sug[i-1]}")
        logger.info(f"审查依据{i}：{audit_foundation[i-1]}")
        logger.info(f"response: {response}")
        Doc.add_paragraph(response)

    Doc.save(target_path)
    return target_path

class PlanAudit:

    def __init__(self, plan_rag = None, client = None, base_info = None):
        self.plan_rag = plan_rag
        self.client = client
        self.plan_rule = PlanRAG(index_path=INDEX_PATH)
        self.agent = PlanAgent(plan_rag, client)
        if base_info is None:
            self.base_info = self.get_base_info()
        else:
            self.base_info = base_info

    def get_answer_from_db(
            self,
            query : Union[str],
            audit_foundation : Union[str],
            topk : Union[int], 
            top_k : Union[int] ) -> Dict[str,str]:
        
        # 导则中针对该问题的具体标准
        database, _ = self.plan_rule.advance_search(query, topk)
        logger.info(f"导则标准:{database}")
        # 报告书中所对应的论证依据
        reports, _ = self.plan_rag.advance_search(query, top_k)
        logger.info(f"论证内容:{reports}")
        logger.info(f"审查要求：{query}")
        if len(reports) == 0:
            # 没检索到对应内容
            return {"input": "根据已知信息无法回答该问题，论证依据不充分",
                    "info": "根据已知信息无法回答该问题，论证依据不充分"
                    }
        else:
            concat_context = ''
            for i in range(top_k):
                concat_context += f"{reports[i]['context']}\n"
            AUDIT_PROPMT_TEMPLATE_FINISH = AUDIT_PROPMT_TEMPLATE1.format(context = concat_context, standard = database, question = query, foundation = audit_foundation)
            return {"input": AUDIT_PROPMT_TEMPLATE_FINISH,
                    "info": concat_context
                    }

    def chat_llm(self, query: str, foundation: str, topk: int, top_k: int, history: List):
        db_response = self.get_answer_from_db(query, foundation, topk, top_k)
        input_text = db_response['input']
        value_info, history = self.client.chat(input_text , history=history)
        value_info = re.sub(r'\n\n(?=[2-9]|[1-9]\d)', '\n', value_info)
        value_info = self.clean_qwen2(value_info)
        logger.info(f"value_info: {value_info}")
        AUDIT_PROPMT_TEMPLATE_FINISH2 = AUDIT_PROPMT_TEMPLATE2.format(base_information = self.base_info, question = query, foundation = foundation, value_info = value_info)
        llm_response, history = self.client.chat(AUDIT_PROPMT_TEMPLATE_FINISH2 , history=history)
        llm_response = re.sub(r'\n\n(?=[2-9]|[1-9]\d)', '\n', llm_response)
        llm_response = self.clean_qwen2(llm_response)
        logger.info(f"llm_response:{llm_response}")
        return llm_response, history

    def get_answer_from_DB(
            self,
            query : Union[str],
            top_k : Union[int] ) -> Dict[str,str]:
        
        # 报告书中所对应的论证依据
        reports, _ = self.plan_rag.advance_search(query, top_k)
        if len(reports) == 0:
            # 没检索到对应内容
            return {"input": "根据已知信息无法回答该问题，论证依据不充分",
                    "info": "根据已知信息无法回答该问题，论证依据不充分"
                    }
        else:
            concat_context = ''
            for i in range(top_k):
                concat_context += f"{reports[i]['context']}\n"
            FORMAT_INFO_PROMPT_FINISH = FORMAT_INFO_PROMPT.format(base_information = concat_context)
            return {"input": FORMAT_INFO_PROMPT_FINISH,
                    "info": concat_context
                    }
        
    def get_base_info(self):
        history = []
        query = "项目名称、申请单位、用海类型、论证工作等级、占用岸线、用海方式、用海面积。"
        db_response = self.get_answer_from_DB(query, 5)
        input_text = db_response['input']
        value_info, history = self.client.chat(input_text , history=history)
        value_info = re.sub(r'\n\n(?=[2-9]|[1-9]\d)', '\n', value_info)
        value_info = self.clean_qwen2(value_info)
        logger.info(f"value_info: {value_info}")
        BASE_INFO_PROMPT_FINISH = BASE_INFO_PROMPT.format(base_information = value_info)
        llm_response, history = self.client.chat(BASE_INFO_PROMPT_FINISH , history=history)
        llm_response = re.sub(r'\n\n(?=[2-9]|[1-9]\d)', '\n', llm_response)
        llm_response = self.clean_qwen2(llm_response)
        llm_response = BeautifulSoup(markdown(llm_response), "html.parser").get_text()
        logger.info(f"llm_response:{llm_response}")
        return llm_response

    def clean_qwen2(self, response: str):
        response = response.rstrip()
        response = response.rstrip('\n')
        if response.endswith("'"):
            response = response[:-1]
            return self.clean_qwen2(response)
        elif response.endswith("')"):
            response = response[:-2]
            return self.clean_qwen2(response)
        elif response.endswith("'("):
            response = response[:-2]
            return self.clean_qwen2(response)
        elif response.endswith(")'"):
            response = response[:-2]
            return self.clean_qwen2(response)
        elif response.endswith("('"):
            response = response[:-2]
            return self.clean_qwen2(response)
        elif response.endswith(")"):
            response = response[:-1]
            return self.clean_qwen2(response)
        elif response.endswith("("):
            response = response[:-1]
            return self.clean_qwen2(response)
        else:
            return response
        return response

if __name__ == '__main__':
    client = ChatQwen()
    report_path = './DATA/reports/raw/深圳市大鹏新区西涌滨海项目海域使用论证报告书报批稿盖章版.pdf'
    filename = os.path.basename(report_path)
    target_path = os.path.join(DB_PATH, filename.split('.')[0])
    plan_rag = PlanRAG(target_path)
    
    filepath = "./DATA/guideline/audit_re.txt"
    filepath2 = "./DATA/guideline/audit_foundation.txt"
    audit_sug = load_txt(filepath)
    audit_foundation = load_txt(filepath2)
    audit_agent = PlanAudit(plan_rag, client)

    # random_indices = random.sample(range(1, 67), 1)

    list = [5, 21, 31, 40]
    for i in list:
        response, history_rs = audit_agent.chat_llm(audit_sug[i-1], audit_foundation[i-1], 1, 2, [])
        logger.info(f"审查要求：{audit_sug[i-1]}")
        logger.info(f"审查依据：{audit_foundation[i-1]}")
        logger.info(f"response: {response}")