
import pdfplumber
import re

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

from loguru import logger
from client import ChatQwen
from PlanRAG import PlanRAG
from config import *
from PlanAgent import PlanCheckAgent
import os
import datetime

def is_chapter1(line):
    pattern1 = r'^\d+(\s)*[^\d\W]+(\s)*(\.)+(\s)+(\d)*$' # 目录中
    return re.match(pattern1, line) is not None

def is_chapter2(line):
    pattern2 = r'^\d+(\s)*[^\d\W]+$' # 正文中
    return re.match(pattern2, line) is not None

def is_attachment1(line):
    pattern1 = r'^[\u4e00-\u9fa5]+[\u4ef6]\d+(\s)*[^\d\W]+(\s)*(\.)+(\s)+(\d)*$' # 目录中
    return re.match(pattern1, line) is not None

def is_attachment2(line):
    pattern2 = r'^[\u4e00-\u9fa5]+[\u4ef6]\d+(\s)*[^\d\W]+$' # 正文中
    return re.match(pattern2, line) is not None

# 检测是否为标题的函数
def is_title(line):
    # TODO 正则表达式匹配类似 "3.8.1.2.3 业务接口" 的格式
    # 确保后面紧跟着的是文字而不是数字或符号
    pattern = r'^\d+(\s)*(\.\d+)+\s+[^\d\W]+$'
    return re.match(pattern, line) is not None

# 检测有图文字的函数 图 1.1-1 xxx
def is_image(line):
    # 确保后面紧跟着的是文字而不是数字或符号
    # ^[\u56fe] 以“图”开头，(\s)*:可能有空格 \d+:一个或多个数字 (\.\d+)*：匹配一个点号后跟一个或多个数字的组合  
    # (\-)*：匹配“-”  [\w\s]+：匹配零个或多个单词字符（字母、数字、下划线）或空白字符  .*：匹配任意数量的任何字符，直到字符串结束。
    pattern = r'^[\u56fe](\s)*\d+(\.\d+)*(\-)*[\w\s]+.*'
    return re.match(pattern, line) is not None

# 检测有表文字的函数 表 1.1-1 xxx
def is_table(line):
    # 确保后面紧跟着的是文字而不是数字或符号(同理图) \u8868->“表”
    pattern = r'^[\u8868](\s)*\d+(\.\d+)*(\-)*[\w\s]+.*'
    return re.match(pattern, line) is not None

# 找到报告提及的图号
def find_image(line):
    # [\u4e00-\u9fa5]+表示一个或多个汉字在“图”的前面，区分is_image
    pattern = r'^[\u4e00-\u9fa5]+[\u56fe](\s)*\d+(\.\d+)*(\-)*[\w\s]+.*'
    return re.match(pattern, line) is not None

# 找到报告提及的表号
def find_table(line):
    # [\u4e00-\u9fa5]+表示一个或多个汉字在“表”的前面，区分is_table
    pattern = r'^[\u4e00-\u9fa5]+[\u8868](\s)*\d+(\.\d+)*(\-)*[\w\s]+.*'
    return re.match(pattern, line) is not None

# 提取筛选内容的函数
def extract_content(pdf_path):
    blank_pages = []
    titles = []
    chapters = []
    chapters1 = []
    chapters2 = []
    attachments = []
    attachments1 = []
    attachments2 = []
    images = []
    tables = []
    m_images = []
    m_tables = []
    with pdfplumber.open(pdf_path) as pdf:
        for page_number, page in enumerate(pdf.pages):
            text = page.extract_text().strip()
            page_images = page.images
            page_tables = page.extract_tables()
            if not text and not page_images and not page_tables:
                # 页面没有文本、图片和表格
                blank_pages.append(page_number + 1)
            else:
                # 如果有内容，进一步处理
                if text:
                    for line in text.split('\n'):
                        if is_title(line):
                            titles.append(line.strip())
                        if is_chapter1(line):
                            chapters1.append(line.strip())
                        if is_chapter2(line):
                            chapters2.append(line.strip())
                        if is_attachment1(line):
                            attachments1.append(line.strip())
                        if is_attachment2(line):
                            attachments2.append(line.strip())
                        if is_image(line):
                            images.append(line.strip())
                        if is_table(line):
                            tables.append(line.strip())
                        if find_image(line):
                            m_images.append(line.strip())
                        if find_table(line):
                            m_tables.append(line.strip())

        # 删除所有空格
        chapters1 = [item.replace(" ", "") for item in chapters1]
        chapters2 = [item.replace(" ", "") for item in chapters2]
        for chapter in chapters2:
            found = False
            for item in chapters1:
                if chapter in item:
                    found = True
                    break
            if found:
                chapters.append(chapter)
        
        # 删除所有空格
        attachments1 = [item.replace(" ", "") for item in attachments1]
        attachments2 = [item.replace(" ", "") for item in attachments2]
        for attachment in attachments2:
            found = False
            for item in attachments1:
                if attachment in item:
                    found = True
                    break
            if found:
                attachments.append(attachment)
    return blank_pages, chapters, titles, attachments, images, tables, m_images, m_tables

def clean_qwen2(response: str):
    response = response.rstrip()
    response = response.rstrip('\n')
    if response.endswith("'"):
        response = response[:-1]
        return clean_qwen2(response)
    elif response.endswith("')"):
        response = response[:-2]
        return clean_qwen2(response)
    elif response.endswith("'("):
        response = response[:-2]
        return clean_qwen2(response)
    elif response.endswith(")'"):
        response = response[:-2]
        return clean_qwen2(response)
    elif response.endswith("('"):
        response = response[:-2]
        return clean_qwen2(response)
    elif response.endswith(")"):
        response = response[:-1]
        return clean_qwen2(response)
    elif response.endswith("("):
        response = response[:-1]
        return clean_qwen2(response)
    else:
        return response
    return response

def extract_sentences(pdf_path):
    with pdfplumber.open(pdf_path) as pdf:
        text = ""
        for page_number, page in enumerate(pdf.pages):
            page_text = page.extract_text().strip()
            text += page_text
        text = text.replace("\n","")

        # 以句号为分隔符划分句子
        sentences = text.split('。')
        # 确保每个句子都以句号结尾
        sentences = [sentence + '。' for sentence in sentences if sentence]

        # 删除首个字符不是汉字的句子
        sentences = [sentence for sentence in sentences if re.match(r'^[\u4e00-\u9fff]', sentence)]
        # 返回句子列表
        return sentences

def mkdoc(pdf_path, filename):
    # 创建一个空白文档
    Doc = Document()
    Doc.styles['Normal'].font.name = u'仿宋'
    Doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
    Doc.styles['Normal'].font.size = Pt(10.5)

    # 添加文档主标题
    Doc.add_heading("合规性审查报告-" + filename.split('.')[0], level=0)

    blank_pages, chapters, titles, attachments, images, tables, m_images, m_tables = extract_content(pdf_path)

    # 1. 空白页检查
    Doc.add_heading("1. 空白页检查", level=1)
    if blank_pages == []:
        Doc.add_paragraph("暂未发现该论证报告含有空白页。")
    else:
        Doc.add_paragraph("该论证报告存在空白页，如下所示：" + ", ".join(map(str, blank_pages)) + "。\n")

    # 初始化变量
    target_path = os.path.join(DB_PATH, filename.split('.')[0])
    plan_rag = PlanRAG(target_path)
    client = ChatQwen()
    plan_check_agent = PlanCheckAgent(plan_rag, client)
    history = []

    # 2. 章节错乱检测
    Doc.add_heading("2. 章节错乱检测（出现漏节、跳节）", level=1)
    CHAPTER_PROMPT_TEMPLATE_FINISH = CHAPTER_PROMPT_TEMPLATE.format(content=chapters)
    llm_response, history = client.chat(CHAPTER_PROMPT_TEMPLATE_FINISH, history=history)
    llm_response = re.sub(r'\n\n(?=[2-9]|[1-9]\d)', '\n', llm_response)
    llm_response = clean_qwen2(llm_response)
    logger.info(f"llm_response_章节错乱检测:{llm_response}")
    Doc.add_paragraph(llm_response)

    # 3. 标题目录级别问题
    Doc.add_heading("3. 标题目录级别弄混、重复", level=1)
    TITLE_PROMPT_TEMPLATE_FINISH = TITLE_PROMPT_TEMPLATE.format(content=titles)
    llm_response, history = client.chat(TITLE_PROMPT_TEMPLATE_FINISH, history=history)
    llm_response = re.sub(r'\n\n(?=[2-9]|[1-9]\d)', '\n', llm_response)
    llm_response = clean_qwen2(llm_response)
    logger.info(f"llm_response_标题目录级别问题:{llm_response}")
    Doc.add_paragraph(llm_response)

    # 4. 错别字和书写问题
    Doc.add_heading("4. 错别字、书写有误（部分，仅测试前20句作为参考）", level=1)
    response = ""
    sentences = extract_sentences(pdf_path)
    count = 0
    for sentence in sentences[1:20]:
        CUOBIEZI_PROMPT_TEMPLATE_FINISH = CUOBIEZI_PROMPT_TEMPLATE.format(content=sentence)
        llm_response, history = client.chat(CUOBIEZI_PROMPT_TEMPLATE_FINISH, history=history)
        llm_response = re.sub(r'\n\n(?=[2-9]|[1-9]\d)', '\n', llm_response)
        llm_response = clean_qwen2(llm_response)
        logger.info(f"llm_response_错别字和书写问题:{llm_response}")
        if "明显" not in llm_response:
            response += llm_response + "\n"
            count += 1

    if count != 0:
        Doc.add_paragraph(response)
    else:
        Doc.add_paragraph("暂未发现有错别字、书写有误。")

    # 5. 报告固定内容缺失
    Doc.add_heading("5. 报告固定内容缺失", level=1)
    ATTACHMENT_PROMPT_TEMPLATE_FINISH = ATTACHMENT_PROMPT_TEMPLATE.format(content=attachments)
    llm_response, history = client.chat(ATTACHMENT_PROMPT_TEMPLATE_FINISH, history=history)
    llm_response = re.sub(r'\n\n(?=[2-9]|[1-9]\d)', '\n', llm_response)
    llm_response = clean_qwen2(llm_response)
    logger.info(f"llm_response_报告固定内容缺失:{llm_response}")
    Doc.add_paragraph(llm_response)

    # 6. 图表序号问题
    Doc.add_heading("6. 图表序号错乱、重复、遗漏等", level=1)
    TUBIAO_PROMPT_TEMPLATE_FINISH = TUBIAO_PROMPT_TEMPLATE.format(image_content=images, table_content=tables)
    llm_response, history = client.chat(TUBIAO_PROMPT_TEMPLATE_FINISH, history=history)
    llm_response = re.sub(r'\n\n(?=[2-9]|[1-9]\d)', '\n', llm_response)
    llm_response = clean_qwen2(llm_response)
    logger.info(f"llm_response_图表序号问题:{llm_response}")
    Doc.add_paragraph(llm_response)

    # 7. 图表引用问题
    Doc.add_heading("7. 报告内容漏写图表号、提及图表号但无对应图表", level=1)
    TUBIAOCHECK_PROMPT_TEMPLATE_FINISH = TUBIAOCHECK_PROMPT_TEMPLATE.format(
        image_mention=m_images, image_content=images, table_mention=m_tables, table_content=tables)
    llm_response, history = client.chat(TUBIAOCHECK_PROMPT_TEMPLATE_FINISH, history=history)
    llm_response = re.sub(r'\n\n(?=[2-9]|[1-9]\d)', '\n', llm_response)
    llm_response = clean_qwen2(llm_response)
    logger.info(f"llm_response_图表引用问题:{llm_response}")

    # 8. 法规合规性检查
    Doc.add_heading("8. 是否符合相关法规和政策要求、是否符合编制要求", level=1)
    law_response = plan_check_agent.check_laws()
    Doc.add_paragraph(law_response)

    # 保存文档
    time_str = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
    target_files = os.path.join(OUT_PATH, filename.split('.')[0])
    if not os.path.exists(target_files):
        os.makedirs(target_files)
    target_path = os.path.join(target_files, f"格式审查综合报告_{time_str}.docx")
    Doc.save(target_path)
    return target_path

def get_responses(pdf_path):
    
    blank_pages, chapters, titles, attachments, images, tables, m_images, m_tables = extract_content(pdf_path)
    
    if blank_pages == []:
        response1 = "1.空白页检查：\n暂未发现该论证报告含有空白页。\n"
    else:
        response1 = "1.空白页检查：\n该论证报告存在空白页，如下所示："+blank_pages+"。"

    # client = OpenAIAPIModel()
    client = ChatQwen()

    history = []
    CHAPTER_PROMPT_TEMPLATE_FINISH = CHAPTER_PROMPT_TEMPLATE.format(content=chapters)
    TITLE_PROMPT_TEMPLATE_FINISH = TITLE_PROMPT_TEMPLATE.format(content=titles)

    ATTACHMENT_PROMPT_TEMPLATE_FINISH = ATTACHMENT_PROMPT_TEMPLATE.format(content=attachments)
    TUBIAO_PROMPT_TEMPLATE_FINISH = TUBIAO_PROMPT_TEMPLATE.format(image_content=images,table_content=tables)
    TUBIAOCHECK_PROMPT_TEMPLATE_FINISH = TUBIAOCHECK_PROMPT_TEMPLATE.format(image_mention=m_images,image_content=images,table_mention=m_tables,table_content=tables)

    llm_response, history = client.chat(CHAPTER_PROMPT_TEMPLATE_FINISH , history=history)
    llm_response = re.sub(r'\n\n(?=[2-9]|[1-9]\d)', '\n', llm_response)
    llm_response = clean_qwen2(llm_response)
    response2 = "2.章节错乱检测（出现漏节、跳节）：\n"+llm_response+"\n"

    llm_response, history = client.chat(TITLE_PROMPT_TEMPLATE_FINISH , history=history)
    llm_response = re.sub(r'\n\n(?=[2-9]|[1-9]\d)', '\n', llm_response)
    llm_response = clean_qwen2(llm_response)
    response3= "3.标题目录级别弄混、重复：\n"+llm_response+"\n"
    
    # 仅显示前20条句子的语法问题
    response = ""
    sentences = extract_sentences(pdf_path)
    for sentence in sentences[1:20]:
        CUOBIEZI_PROMPT_TEMPLATE_FINISH = CUOBIEZI_PROMPT_TEMPLATE.format(content=sentence)
        llm_response, history = client.chat(CUOBIEZI_PROMPT_TEMPLATE_FINISH , history=history)
        llm_response = re.sub(r'\n\n(?=[2-9]|[1-9]\d)', '\n', llm_response)
        llm_response = clean_qwen2(llm_response)
        if "明显" not in llm_response:
            response += llm_response + "\n"
    
    response4 = "4.错别字、书写有误：\n"+response+"\n"

    llm_response, history = client.chat(ATTACHMENT_PROMPT_TEMPLATE_FINISH , history=history)
    llm_response = re.sub(r'\n\n(?=[2-9]|[1-9]\d)', '\n', llm_response)
    llm_response = clean_qwen2(llm_response)
    response5 = "5.报告固定内容缺失：\n"+llm_response+"\n"

    llm_response, history = client.chat(TUBIAO_PROMPT_TEMPLATE_FINISH , history=history)
    llm_response = re.sub(r'\n\n(?=[2-9]|[1-9]\d)', '\n', llm_response)
    llm_response = clean_qwen2(llm_response)
    response6 = "6.图表序号错乱、重复、遗漏等: \n" +llm_response + "\n"

    llm_response, history = client.chat(TUBIAOCHECK_PROMPT_TEMPLATE_FINISH , history=history)
    llm_response = re.sub(r'\n\n(?=[2-9]|[1-9]\d)', '\n', llm_response)
    llm_response = clean_qwen2(llm_response)
    response7 = "7.报告内容漏写图表号、报告中提及图表号但无对应图表: \n" +llm_response + "\n"

    return response1, response2, response3, response4, response5, response6, response7

if __name__ == '__main__' :
    report_path = './DATA/reports/raw/test.pdf'
    filename = os.path.basename(report_path)
    target_path = os.path.join(DB_PATH, filename.split('.')[0])
    plan_rag = PlanRAG(target_path)
    client = ChatQwen()
    agent = PlanCheckAgent(plan_rag, client)
    response = agent.check_laws()
    logger.info(f"response:{response}")
    mkdoc(report_path, response)
    # sentences = extract_sentences(path_file)
    # client = OpenAIAPIModel()
    # client = ChatQwen()
    # response = []
    # history = []
    # for sentence in sentences[1:10]:
    #     CUOBIEZI_PROMPT_TEMPLATE_FINISH = CUOBIEZI_PROMPT_TEMPLATE.format(content=sentence)
    #     llm_response, history = client.chat(CUOBIEZI_PROMPT_TEMPLATE_FINISH , history=history)
    #     llm_response = re.sub(r'\n\n(?=[2-9]|[1-9]\d)', '\n', llm_response)
    #     logger.info(f"response: {llm_response}")
    #     response += llm_response + "\n"

    # with open(path_file, 'r', encoding='utf-8') as f:
    #     content = f.read()
    # paragraphs = re.findall(r'^[^#].*$', content, re.MULTILINE)
    # for paragraph in paragraphs:
    #     CUOBIEZI_PROMPT_TEMPLATE_FINISH = CUOBIEZI_PROMPT_TEMPLATE.format(content=paragraph)
    #     llm_response, history = client.chat(CUOBIEZI_PROMPT_TEMPLATE_FINISH , history=history)
    #     llm_response = re.sub(r'\n\n(?=[2-9]|[1-9]\d)', '\n', llm_response)
    #     logger.info(f"llm_response: {llm_response}")
    #     response += llm_response